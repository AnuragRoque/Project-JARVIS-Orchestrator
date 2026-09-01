"""Extract readable text from a file: plain text, PDF, or Word (.docx)."""
from __future__ import annotations

import os

from jarvis.app.logsetup import get_logger

log = get_logger("documents.reader")

_TEXT_EXTS = {
    ".txt", ".md", ".markdown", ".rst", ".log", ".csv", ".tsv", ".json",
    ".xml", ".yaml", ".yml", ".ini", ".cfg", ".toml", ".env",
    ".py", ".js", ".ts", ".tsx", ".jsx", ".html", ".htm", ".css", ".java",
    ".c", ".cpp", ".h", ".cs", ".go", ".rs", ".rb", ".php", ".sh", ".ps1",
    ".bat", ".sql",
}
_MAX_PDF_PAGES = 50


def read_document(path: str, max_chars: int = 8000) -> dict:
    """Return {ok, kind, text, ...} for a file, or {ok: False, error}."""
    if not path or not os.path.isfile(path):
        return {"ok": False, "error": f"File not found: {path}"}

    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            text, meta = _read_pdf(path)
            kind = "pdf"
        elif ext == ".docx":
            text, meta = _read_docx(path)
            kind = "docx"
        elif ext in _TEXT_EXTS or _looks_textual(path):
            text, meta = _read_text(path), {}
            kind = "text"
        else:
            return {"ok": False,
                    "error": f"I can't read '{ext or 'this'}' files as text yet "
                             f"(supported: text, .pdf, .docx)."}
    except Exception as exc:
        log.exception("read_document failed for %s", path)
        return {"ok": False, "error": f"Couldn't read the file: {exc}"}

    text = text or ""
    full_len = len(text)
    truncated = full_len > max_chars
    result = {
        "ok": True,
        "path": path,
        "name": os.path.basename(path),
        "kind": kind,
        "chars": full_len,
        "truncated": truncated,
        "text": text[:max_chars] + ("\n…[truncated]" if truncated else ""),
    }
    result.update(meta)
    if not text.strip():
        result["note"] = ("No extractable text — the file may be empty or a "
                          "scanned/image-only PDF (OCR isn't available yet).")
    return result


# ------------------------------------------------------------------ readers
def _read_text(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as fh:
        return fh.read()


def _read_pdf(path: str) -> tuple[str, dict]:
    from pypdf import PdfReader
    reader = PdfReader(path)
    pages = reader.pages
    chunks = []
    for i, page in enumerate(pages[:_MAX_PDF_PAGES]):
        try:
            chunks.append(page.extract_text() or "")
        except Exception:
            chunks.append("")
    meta = {"pages": len(pages)}
    if len(pages) > _MAX_PDF_PAGES:
        meta["pages_read"] = _MAX_PDF_PAGES
    return "\n\n".join(c for c in chunks if c), meta


def _read_docx(path: str) -> tuple[str, dict]:
    import docx
    d = docx.Document(path)
    paras = [p.text for p in d.paragraphs]
    # Include simple table text too.
    for table in d.tables:
        for row in table.rows:
            paras.append(" | ".join(c.text for c in row.cells))
    return "\n".join(p for p in paras if p is not None), {}


def _looks_textual(path: str, sniff: int = 2048) -> bool:
    """Heuristic: treat a file with no NUL bytes and mostly printable content as text."""
    try:
        with open(path, "rb") as fh:
            chunk = fh.read(sniff)
    except OSError:
        return False
    if not chunk or b"\x00" in chunk:
        return False
    printable = sum(1 for b in chunk if 9 <= b <= 13 or 32 <= b <= 126 or b >= 128)
    return printable / len(chunk) > 0.85
